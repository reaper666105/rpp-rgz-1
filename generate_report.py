"""
Скрипт для генерации docx отчёта по РГЗ (вариант 12).
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def setup_document_style(doc: Document):
    """Настройка стилей документа согласно требованиям РГЗ."""
    # Основной стиль для текста
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(21)  # 1.5 * 14pt
    style.paragraph_format.first_line_indent = Inches(0.5)  # ~1.25 см
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)


def add_title_page(doc: Document):
    """Титульный лист."""
    # Пустая страница для титульного листа (без номера)
    para = doc.add_paragraph()
    para.add_run().add_break()
    
    # Название университета/факультета
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('Факультет бизнеса\nКафедра экономической информатики')
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Пустые строки
    for _ in range(8):
        doc.add_paragraph()
    
    # Название работы
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('РАСЧЕТНО-ГРАФИЧЕСКОЕ ЗАДАНИЕ')
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('по дисциплине «Разработка программных приложений»')
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('Вариант 12')
    run.font.size = Pt(14)
    run.font.bold = True
    
    for _ in range(10):
        doc.add_paragraph()
    
    # Студент
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run('Выполнил(а): _________________')
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run('Проверил(а): _________________')
    run.font.size = Pt(14)
    
    for _ in range(3):
        doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('Новосибирск, 2024')
    run.font.size = Pt(14)


def add_heading_style(doc: Document, level: int, text: str, center: bool = False):
    """Добавление заголовка с правильным форматированием."""
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    
    run = para.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    
    return para


def add_regular_paragraph(doc: Document, text: str):
    """Добавление обычного абзаца с правильным форматированием."""
    para = doc.add_paragraph(text)
    para.paragraph_format.first_line_indent = Inches(0.5)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = Pt(21)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    return para


def add_code_block(doc: Document, code: str, language: str = ""):
    """Добавление блока кода."""
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    return para


def setup_page_numbers(doc: Document):
    """Настройка нумерации страниц справа внизу."""
    for section in doc.sections:
        footer = section.footer
        
        # Очищаем существующие параграфы
        for para in footer.paragraphs:
            para.clear()
        
        # Если параграфов нет, создаём новый
        if not footer.paragraphs:
            para = footer.add_paragraph()
        else:
            para = footer.paragraphs[0]
        
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = para.add_run()
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        
        # Добавляем поле номера страницы
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)


def read_source_file(filepath: str) -> str:
    """Чтение исходного файла для приложения."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"# Ошибка чтения файла {filepath}: {e}"


def main():
    doc = Document()
    setup_document_style(doc)
    
    # Титульный лист (без номера)
    add_title_page(doc)
    
    # Разрыв страницы для содержания
    doc.add_page_break()
    
    # Содержание
    add_heading_style(doc, 1, 'СОДЕРЖАНИЕ', center=True)
    
    toc_items = [
        ('Введение', 3),
        ('Глава 1. Описание используемых технологий', 4),
        ('Глава 2. Описание выполненной работы', 5),
        ('Заключение', 6),
        ('Список использованных источников', 7),
        ('Приложение', 8),
    ]
    
    for item, page in toc_items:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        
        run1 = para.add_run(item)
        run1.font.size = Pt(14)
        run1.font.name = 'Times New Roman'
        
        # Табуляция для номеров страниц
        para.add_run('\t' * 5)
        
        run2 = para.add_run(str(page))
        run2.font.size = Pt(14)
        run2.font.name = 'Times New Roman'
    
    # Разрыв страницы
    doc.add_page_break()
    
    # ВВЕДЕНИЕ
    add_heading_style(doc, 1, 'ВВЕДЕНИЕ', center=True)
    
    intro_text = (
        "В рамках выполнения расчетно-графического задания по дисциплине "
        "«Разработка программных приложений» было разработано RESTful API "
        "на Flask для управления запасами товаров (инвентарем) на складе. "
        "Приложение поддерживает операции по добавлению, обновлению и удалению "
        "товаров, а также автоматическую генерацию отчетов о состоянии инвентаря. "
        "Для обеспечения качества кода и автоматизации процессов разработки "
        "настроен workflow для CI/CD в GitHub Actions."
    )
    add_regular_paragraph(doc, intro_text)
    
    doc.add_page_break()
    
    # ГЛАВА 1
    add_heading_style(doc, 1, 'ГЛАВА 1. ОПИСАНИЕ ИСПОЛЬЗУЕМЫХ ТЕХНОЛОГИЙ', center=True)
    
    add_heading_style(doc, 2, '1.1. Flask')
    text1_1 = (
        "Flask — это легковесный веб-фреймворк для Python, разработанный "
        "Армином Ронахером. Flask предоставляет минималистичный подход к "
        "созданию веб-приложений, позволяя разработчику самостоятельно выбирать "
        "инструменты и библиотеки для решения конкретных задач. В данной работе "
        "использована версия Flask 3.0.3."
    )
    add_regular_paragraph(doc, text1_1)
    
    text1_1_2 = (
        "Основные преимущества Flask включают простоту использования, гибкость "
        "архитектуры, наличие большого количества расширений и активное сообщество. "
        "Flask идеально подходит для создания RESTful API благодаря поддержке "
        "маршрутизации, обработки HTTP-запросов и интеграции с различными "
        "библиотеками для работы с базами данных."
    )
    add_regular_paragraph(doc, text1_1_2)
    
    add_heading_style(doc, 2, '1.2. PostgreSQL и SQLAlchemy')
    text1_2 = (
        "PostgreSQL — это мощная объектно-реляционная система управления базами "
        "данных с открытым исходным кодом. PostgreSQL обеспечивает высокую "
        "надежность, целостность данных и поддержку сложных запросов. В проекте "
        "используется PostgreSQL версии 16 для хранения информации о товарах на складе."
    )
    add_regular_paragraph(doc, text1_2)
    
    text1_2_2 = (
        "SQLAlchemy — это библиотека Python для работы с базами данных, которая "
        "предоставляет инструменты объектно-реляционного отображения (ORM). "
        "SQLAlchemy версии 2.0.36 позволяет работать с базой данных на высоком "
        "уровне абстракции, используя Python-объекты вместо прямых SQL-запросов. "
        "Flask-SQLAlchemy версии 3.1.1 обеспечивает интеграцию SQLAlchemy с Flask, "
        "упрощая работу с базой данных в веб-приложении."
    )
    add_regular_paragraph(doc, text1_2_2)
    
    add_heading_style(doc, 2, '1.3. Pytest')
    text1_3 = (
        "Pytest — это популярный фреймворк для написания и запуска тестов в Python. "
        "Pytest версии 8.3.4 используется в проекте для написания unit-тестов, "
        "проверяющих корректность работы всех эндпоинтов API. Pytest предоставляет "
        "удобный синтаксис для написания тестов, автоматическое обнаружение тестовых "
        "файлов и подробные отчеты о результатах выполнения тестов."
    )
    add_regular_paragraph(doc, text1_3)
    
    add_heading_style(doc, 2, '1.4. Bandit')
    text1_4 = (
        "Bandit — это инструмент для статического анализа безопасности Python-кода. "
        "Bandit версии 1.8.0 используется в проекте для выявления потенциальных "
        "уязвимостей и проблем безопасности. Инструмент интегрирован в CI/CD pipeline "
        "GitHub Actions и автоматически проверяет код при каждом push в основную ветку."
    )
    add_regular_paragraph(doc, text1_4)
    
    add_heading_style(doc, 2, '1.5. GitHub Actions')
    text1_5 = (
        "GitHub Actions — это платформа для автоматизации процессов разработки "
        "программного обеспечения, интегрированная в GitHub. GitHub Actions позволяет "
        "создавать workflows (рабочие процессы) для автоматического выполнения "
        "различных задач, таких как запуск тестов, проверка кода и развертывание "
        "приложений. В проекте настроен workflow для автоматического запуска "
        "unit-тестов и проверки безопасности кода при каждом push в ветку main."
    )
    add_regular_paragraph(doc, text1_5)
    
    add_heading_style(doc, 2, '1.6. Дополнительные библиотеки')
    text1_6 = (
        "В проекте также используются следующие библиотеки: python-dotenv (версия 1.0.1) "
        "для загрузки переменных окружения из файла .env, psycopg2-binary (версия 2.9.10) "
        "как адаптер для подключения к PostgreSQL из Python."
    )
    add_regular_paragraph(doc, text1_6)
    
    doc.add_page_break()
    
    # ГЛАВА 2
    add_heading_style(doc, 1, 'ГЛАВА 2. ОПИСАНИЕ ВЫПОЛНЕННОЙ РАБОТЫ', center=True)
    
    add_heading_style(doc, 2, '2.1. Архитектура приложения')
    text2_1 = (
        "Приложение разработано с использованием архитектурного паттерна Application Factory, "
        "который позволяет создавать экземпляры Flask-приложения с различными конфигурациями. "
        "Проект организован в виде модульной структуры: модуль app содержит основную логику "
        "приложения, модуль tests — unit-тесты, файл wsgi.py служит точкой входа для запуска приложения."
    )
    add_regular_paragraph(doc, text2_1)
    
    add_heading_style(doc, 2, '2.2. Модель данных')
    text2_2 = (
        "Для хранения информации о товарах создана модель Item с следующими полями: "
        "id (первичный ключ), name (название товара, строка до 200 символов), "
        "quantity (количество, целое число), price (цена за единицу, десятичное число "
        "с точностью до 2 знаков после запятой), category (категория товара, строка "
        "до 100 символов с индексом для ускорения поиска), created_at и updated_at "
        "(временные метки создания и обновления записи)."
    )
    add_regular_paragraph(doc, text2_2)
    
    add_heading_style(doc, 2, '2.3. RESTful API эндпоинты')
    
    add_heading_style(doc, 3, '2.3.1. Управление товарами')
    text2_3_1 = (
        "Реализованы следующие эндпоинты для работы с товарами:\n"
        "• POST /items — добавление нового товара. Принимает JSON-объект с полями name, "
        "quantity, price, category. Выполняется валидация данных: количество не может "
        "быть отрицательным, цена должна быть больше нуля, все поля обязательны.\n"
        "• GET /items — получение списка всех товаров с возможностью фильтрации по категории "
        "через query-параметр category.\n"
        "• GET /items/<id> — получение информации о конкретном товаре по его идентификатору.\n"
        "• PUT /items/<id> — обновление информации о товаре. Позволяет обновить любое из полей: "
        "name, quantity, price, category.\n"
        "• DELETE /items/<id> — удаление товара из базы данных."
    )
    add_regular_paragraph(doc, text2_3_1)
    
    add_heading_style(doc, 3, '2.3.2. Генерация отчетов')
    text2_3_2 = (
        "Эндпоинт GET /reports/summary генерирует сводный отчет о состоянии инвентаря. "
        "Отчет включает:\n"
        "• Общую стоимость всех товаров (сумма quantity * price для всех товаров).\n"
        "• Разбивку по категориям: количество товаров в каждой категории, общее количество "
        "единиц товара в категории, общая стоимость товаров в категории.\n"
        "• Список товаров с нулевым или отрицательным количеством.\n\n"
        "Отчет может быть получен в двух форматах: JSON (по умолчанию) и CSV (при указании "
        "query-параметра format=csv)."
    )
    add_regular_paragraph(doc, text2_3_2)
    
    add_heading_style(doc, 2, '2.4. Валидация данных')
    text2_4 = (
        "При добавлении и обновлении товаров выполняется строгая валидация входных данных:\n"
        "• Поле name должно быть непустой строкой.\n"
        "• Поле quantity должно быть целым неотрицательным числом (quantity >= 0).\n"
        "• Поле price должно быть положительным числом (price > 0).\n"
        "• Поле category должно быть непустой строкой.\n\n"
        "При нарушении правил валидации API возвращает HTTP-статус 400 с описанием ошибки."
    )
    add_regular_paragraph(doc, text2_4)
    
    add_heading_style(doc, 2, '2.5. Unit-тесты')
    text2_5 = (
        "Разработаны unit-тесты, покрывающие все основные сценарии использования API:\n"
        "• Создание товара и получение его по идентификатору.\n"
        "• Валидация данных при создании товара (проверка на отрицательное количество).\n"
        "• Получение списка товаров и фильтрация по категории.\n"
        "• Обновление товара и его удаление.\n"
        "• Генерация сводного отчета в форматах JSON и CSV.\n\n"
        "Тесты используют тестовую базу данных (SQLite для локального запуска, PostgreSQL "
        "в CI/CD pipeline) и изолированы друг от друга благодаря использованию фикстур pytest."
    )
    add_regular_paragraph(doc, text2_5)
    
    add_heading_style(doc, 2, '2.6. Настройка CI/CD в GitHub Actions')
    text2_6 = (
        "Настроен автоматический workflow в GitHub Actions, который запускается при каждом "
        "push в ветку main. Workflow выполняет следующие шаги:\n"
        "• Проверка кода из репозитория.\n"
        "• Установка Python 3.11 и зависимостей проекта.\n"
        "• Запуск PostgreSQL в качестве сервиса для тестирования.\n"
        "• Ожидание готовности базы данных.\n"
        "• Проверка безопасности кода с помощью Bandit.\n"
        "• Запуск unit-тестов с помощью pytest.\n\n"
        "Все шаги должны успешно завершиться, иначе workflow помечается как неуспешный."
    )
    add_regular_paragraph(doc, text2_6)
    
    doc.add_page_break()
    
    # ЗАКЛЮЧЕНИЕ
    add_heading_style(doc, 1, 'ЗАКЛЮЧЕНИЕ', center=True)
    
    conclusion_text = (
        "В рамках выполнения расчетно-графического задания было разработано RESTful API "
        "на Flask для управления запасами товаров на складе. Приложение успешно реализует "
        "все требуемые функции: добавление, обновление, удаление и просмотр товаров, а также "
        "генерацию отчетов в форматах JSON и CSV.\n\n"
        "Приложение использует PostgreSQL для хранения данных, что обеспечивает надежность "
        "и масштабируемость решения. Реализована валидация входных данных и обработка ошибок. "
        "Написаны unit-тесты, покрывающие основные сценарии использования API.\n\n"
        "Настроен автоматический CI/CD pipeline в GitHub Actions, который обеспечивает "
        "проверку качества кода и его безопасности при каждом изменении в репозитории. Это "
        "позволяет выявлять проблемы на ранних этапах разработки и поддерживать высокое "
        "качество кодовой базы."
    )
    add_regular_paragraph(doc, conclusion_text)
    
    doc.add_page_break()
    
    # СПИСОК ИСТОЧНИКОВ (ГОСТ 2021)
    add_heading_style(doc, 1, 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', center=True)
    
    sources = [
        "Flask Documentation [Электронный ресурс]. — Режим доступа: https://flask.palletsprojects.com/ (дата обращения: 17.12.2024).",
        "PostgreSQL: The World's Most Advanced Open Source Relational Database [Электронный ресурс]. — Режим доступа: https://www.postgresql.org/ (дата обращения: 17.12.2024).",
        "SQLAlchemy Documentation [Электронный ресурс]. — Режим доступа: https://docs.sqlalchemy.org/ (дата обращения: 17.12.2024).",
        "Pytest Documentation [Электронный ресурс]. — Режим доступа: https://docs.pytest.org/ (дата обращения: 17.12.2024).",
        "Bandit: A security linter for Python code [Электронный ресурс]. — Режим доступа: https://bandit.readthedocs.io/ (дата обращения: 17.12.2024).",
        "GitHub Actions Documentation [Электронный ресурс]. — Режим доступа: https://docs.github.com/en/actions (дата обращения: 17.12.2024).",
    ]
    
    for i, source in enumerate(sources, 1):
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.hanging_indent = Inches(-0.5)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        
        run = para.add_run(f"{i}. {source}")
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # ПРИЛОЖЕНИЕ
    add_heading_style(doc, 1, 'ПРИЛОЖЕНИЕ', center=True)
    add_heading_style(doc, 2, 'Исходный код приложения')
    
    # Добавляем основные файлы исходного кода
    files_to_include = [
        ('app/__init__.py', 'Файл инициализации приложения'),
        ('app/models.py', 'Модель данных Item'),
        ('app/api.py', 'REST API эндпоинты'),
        ('app/extensions.py', 'Расширения Flask'),
        ('wsgi.py', 'Точка входа приложения'),
        ('tests/test_inventory_api.py', 'Unit-тесты'),
        ('tests/conftest.py', 'Конфигурация pytest'),
        ('.github/workflows/ci.yml', 'GitHub Actions workflow'),
    ]
    
    for filepath, description in files_to_include:
        if os.path.exists(filepath):
            add_heading_style(doc, 3, f'{description} ({filepath})')
            code = read_source_file(filepath)
            add_code_block(doc, code)
            doc.add_paragraph()  # Пустая строка между файлами
    
    # Настройка нумерации страниц
    setup_page_numbers(doc)
    
    # Сохранение документа
    output_path = 'РГЗ_Вариант12_Отчет.docx'
    doc.save(output_path)
    print(f"Отчёт сохранён в файл: {output_path}")


if __name__ == '__main__':
    main()

